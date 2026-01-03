"""
Document service layer
"""
import os
import uuid
import logging
import mimetypes
from datetime import datetime, timezone
from typing import Optional, Tuple, List, Dict, Any
from pathlib import Path
from sqlalchemy.orm import Session
from fastapi import UploadFile, HTTPException, status

from src.core.config import settings
from src.utils.helpers import human_readable_size
from src.documents.models import Document, DocumentStatus
from src.documents.schemas import DocumentUpload
from src.storage.local import LocalStorage
from src.ai.services import ai_processor

logger = logging.getLogger(__name__)


class DocumentService:
    """Document management service"""
    
    def __init__(self):
        self.storage = LocalStorage()
        self.ai_processor = ai_processor
        self.allowed_mime_types = {
            'application/pdf': '.pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
            'text/plain': '.txt',
        }
    
    async def upload_document(
        self,
        db: Session,
        file: UploadFile,
        user_id: int,
        document_metadata: Optional[DocumentUpload] = None,
    ) -> Document:
        """
        Upload and process a document
        """
        # Validate file size
        max_size = settings.MAX_FILE_SIZE_MB * 1024 * 1024
        
        # Read file content for validation
        content = await file.read()
        file_size = len(content)
        
        if file_size > max_size:
            raise HTTPException(
                status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                detail=f"File size exceeds maximum allowed size of {settings.MAX_FILE_SIZE_MB}MB",
            )
        
        # Reset file pointer for storage
        await file.seek(0)
        
        # Validate MIME type
        mime_type = file.content_type
        if mime_type not in self.allowed_mime_types:
            # Also check by file extension as fallback
            file_extension = Path(file.filename).suffix.lower()
            allowed_extensions = list(self.allowed_mime_types.values())
            
            if file_extension not in allowed_extensions:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"File type not allowed. Allowed types: {', '.join(allowed_extensions)}",
                )
            else:
                # Map extension back to MIME type
                extension_to_mime = {v: k for k, v in self.allowed_mime_types.items()}
                mime_type = extension_to_mime.get(file_extension, mime_type)
        
        # Generate unique filename
        file_extension = self.allowed_mime_types.get(mime_type, Path(file.filename).suffix)
        unique_filename = f"{uuid.uuid4()}{file_extension}"
        
        # Save file
        file_path = await self.storage.save_file(
            content=content,
            filename=unique_filename,
            user_id=user_id,
        )
        
        # Create document record
        document = Document(
            filename=unique_filename,
            original_filename=file.filename,
            file_path=file_path,
            file_type=file_extension,
            file_size=file_size,
            mime_type=mime_type,
            user_id=user_id,
            status=DocumentStatus.PENDING,
            document_metadata=document_metadata.dict() if document_metadata else {},
        )
        
        db.add(document)
        db.commit()
        db.refresh(document)
        
        logger.info(f"Document uploaded: {document.id} by user {user_id}")
        
        # Start async processing
        await self._process_document_async(db, document.id)
        
        return document
    
    async def _process_document_async(self, db: Session, document_id: int):
        """
        Process document asynchronously (AI analysis, embeddings, etc.)
        """
        document = None
        try:
            document = db.query(Document).filter(Document.id == document_id).first()
            if not document:
                logger.error(f"Document {document_id} not found")
                return
            
            # Update status
            document.status = DocumentStatus.PROCESSING
            document.processing_started_at = datetime.now(timezone.utc)
            db.commit()
            
            logger.info(f"Processing document {document_id}: {document.original_filename}")
            
            # Read file content
            content = await self.storage.read_file(document.file_path)
            logger.info(f"Read {len(content)} bytes")
            
            # Extract text
            extracted_text = ""
            if document.file_type == '.pdf':
                extracted_text = await self._extract_pdf_text(content)
            elif document.file_type == '.docx':
                extracted_text = await self._extract_docx_text(content)
            elif document.file_type == '.txt':
                extracted_text = content.decode('utf-8')
            
            logger.info(f"Extracted {len(extracted_text)} characters")
            document.extracted_text = extracted_text
            
            # Perform AI analysis
            logger.info("Starting AI analysis...")
            try:
                analysis = await self.ai_processor.analyze_document(extracted_text)
                logger.info(f"AI analysis complete: {analysis.keys()}")
                
                if 'error' in analysis:
                    logger.warning(f"AI analysis had error: {analysis['error']}")
                    document.ai_analysis = analysis
                else:
                    document.ai_analysis = analysis
                    if analysis.get('summary'):
                        document.summary = analysis['summary']
                        logger.info(f"Generated summary: {analysis['summary'][:100]}...")
            except Exception as e:
                logger.error(f"AI analysis failed: {str(e)}", exc_info=True)
                document.ai_analysis = {"error": str(e)}
            
            # Store embeddings (optional - don't fail if this doesn't work)
            try:
                logger.info("Storing embeddings...")
                await self.ai_processor.store_document_embeddings(
                    document_id=document.id,
                    text=extracted_text,
                    metadata={
                        'user_id': document.user_id,
                        'filename': document.original_filename,
                    }
                )
                logger.info("Embeddings stored successfully")
            except Exception as e:
                logger.warning(f"Embedding storage failed (continuing anyway): {str(e)}")
            
            # Update status
            document.status = DocumentStatus.PROCESSED
            document.processing_completed_at = datetime.now(timezone.utc)
            db.commit()
            
            logger.info(f"Document {document_id} processed successfully")
            
        except Exception as e:
            logger.error(f"Error processing document {document_id}: {str(e)}", exc_info=True)
            
            if document:
                document.status = DocumentStatus.FAILED
                document.processing_error = str(e)
                db.commit()
    
    async def _extract_pdf_text(self, content: bytes) -> str:
        """Extract text from PDF with multiple fallback methods"""
        from io import BytesIO
        
        pdf_file = BytesIO(content)
        
        # Method 1: Try PyMuPDF (fitz) - usually the best
        try:
            import fitz  # PyMuPDF
            pdf_file.seek(0)
            doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            
            if len(text.strip()) > 100:
                logger.info(f"PyMuPDF extracted {len(text)} characters")
                return text.strip()
        except ImportError:
            logger.warning("PyMuPDF not installed")
        except Exception as e:
            logger.warning(f"PyMuPDF extraction failed: {e}")
        
        # Method 2: Try pdfplumber
        try:
            import pdfplumber
            pdf_file.seek(0)
            with pdfplumber.open(pdf_file) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if len(text.strip()) > 100:
                logger.info(f"pdfplumber extracted {len(text)} characters")
                return text.strip()
        except ImportError:
            logger.warning("pdfplumber not installed")
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 3: Try PyPDF2 as last resort
        try:
            import PyPDF2
            pdf_file.seek(0)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            
            logger.info(f"PyPDF2 extracted {len(text)} characters")
            return text.strip()
            
        except Exception as e:
            logger.error(f"All PDF extraction methods failed: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract text from PDF. The PDF might be scanned or image-based. Error: {str(e)}",
            )
    
    async def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX"""
        try:
            from docx import Document as DocxDocument
            from io import BytesIO
            
            docx_file = BytesIO(content)
            doc = DocxDocument(docx_file)
            
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return text.strip()
            
        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract text from DOCX: {str(e)}",
            )
    
    def get_user_documents(
        self,
        db: Session,
        user_id: int,
        skip: int = 0,
        limit: int = 100,
        status: Optional[str] = None,
    ) -> Tuple[List[Document], int]:
        """
        Get documents for a user
        """
        query = db.query(Document).filter(Document.user_id == user_id)
        
        if status:
            # Convert string status to DocumentStatus enum
            status_upper = status.upper()
            if status_upper == "PROCESSED":
                query = query.filter(Document.status == DocumentStatus.PROCESSED)
            elif status_upper == "PROCESSING":
                query = query.filter(Document.status == DocumentStatus.PROCESSING)
            elif status_upper == "PENDING":
                query = query.filter(Document.status == DocumentStatus.PENDING)
            elif status_upper == "FAILED":
                query = query.filter(Document.status == DocumentStatus.FAILED)
            else:
                # Try direct string comparison as fallback
                query = query.filter(Document.status == status)
        
        total = query.count()
        documents = query.order_by(
            Document.created_at.desc()
        ).offset(skip).limit(limit).all()
        
        return documents, total
    
    def get_document(
        self,
        db: Session,
        document_id: int,
        user_id: int,
    ) -> Document:
        """
        Get document by ID (with ownership check)
        """
        document = db.query(Document).filter(
            Document.id == document_id,
            Document.user_id == user_id,
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found",
            )
        
        return document
    
    async def delete_document(
        self,
        db: Session,
        document_id: int,
        user_id: int,
    ) -> bool:
        """
        Delete document
        """
        document = self.get_document(db, document_id, user_id)
        
        try:
            # Delete from storage
            await self.storage.delete_file(document.file_path)
            
            # Delete from database
            db.delete(document)
            db.commit()
            
            # Delete embeddings from Pinecone
            await self.ai_processor.delete_document_embeddings(document_id)
            
            logger.info(f"Document deleted: {document_id} by user {user_id}")
            return True
            
        except Exception as e:
            logger.error(f"Error deleting document {document_id}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to delete document",
            )
    
    async def download_document(
        self,
        db: Session,
        document_id: int,
        user_id: int,
    ) -> Tuple[bytes, str]:
        """
        Download document
        """
        document = self.get_document(db, document_id, user_id)
        
        try:
            content = await self.storage.read_file(document.file_path)
            return content, document.original_filename
            
        except Exception as e:
            logger.error(f"Error reading document {document_id}: {str(e)}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to download document",
            )
    
    def get_document_analysis(
        self,
        db: Session,
        document_id: int,
        user_id: int,
    ) -> Dict[str, Any]:
        """
        Get document AI analysis
        """
        document = self.get_document(db, document_id, user_id)
        
        if not document.ai_analysis:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document analysis not available",
            )
        
        return document.ai_analysis


document_service = DocumentService()